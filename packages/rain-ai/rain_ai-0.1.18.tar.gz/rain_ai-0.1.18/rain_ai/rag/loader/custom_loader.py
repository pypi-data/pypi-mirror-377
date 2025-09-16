import os
from pathlib import Path
from typing import Iterator, Literal, Any

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from openpyxl import load_workbook


def _has_page_break(paragraph) -> bool:
    # 检查段落文本中的分页符字符
    if "\f" in paragraph.text or "\x0c" in paragraph.text:
        return True

    # 使用公开的API检查分页符
    for run in paragraph.runs:
        if "\f" in run.text or "\x0c" in run.text:
            return True

    return False


class WordLoader(BaseLoader):
    def __init__(self, file_path: str | Path, mode: Literal["page", "single"] = "page"):
        self.file_path = str(file_path)
        self.mode = mode

        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        if not os.path.isfile(self.file_path):
            raise ValueError(f"文件不存在: {self.file_path}")

    def lazy_load(self) -> Iterator[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(self.file_path)

        if self.mode == "single":
            yield from self._load_as_single(doc)
        else:  # mode == "page"
            yield from self._load_by_pages(doc)

    def _load_as_single(self, doc) -> Iterator[Document]:
        content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        yield Document(
            page_content=content, metadata={"source": self.file_path, "mode": "single"}
        )

    def _load_by_pages(self, doc) -> Iterator[Document]:
        current_page = 1
        page_content = []

        for paragraph in doc.paragraphs:
            # 检查分页符
            if _has_page_break(paragraph):
                if page_content:
                    yield self._create_page_document(page_content, current_page)
                    page_content = []
                    current_page += 1

            if paragraph.text.strip():
                page_content.append(paragraph.text)

        # 输出最后一页
        if page_content:
            yield self._create_page_document(page_content, current_page)

    def _create_page_document(self, content: list, page_num: int) -> Document:
        return Document(
            page_content="\n".join(content),
            metadata={"source": self.file_path, "page": page_num, "mode": "page"},
        )


def _detect_header(worksheet) -> bool:
    if worksheet.max_row < 2:
        return False

    first_row = list(worksheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]

    # 如果第一行主要是文本，可能是表头
    text_cells = 0
    total_cells = 0

    for cell in first_row:
        if cell is not None:
            total_cells += 1
            if isinstance(cell, str) and not cell.isdigit():
                text_cells += 1

    return text_cells > total_cells * 0.5 if total_cells > 0 else False


def _extract_sheet_data(worksheet) -> dict:
    rows = []
    max_row = worksheet.max_row
    max_column = worksheet.max_column

    # 检查是否有数据
    if max_row == 1 and max_column == 1:
        cell_value = worksheet.cell(1, 1).value
        if cell_value is None:
            return {"content": "", "max_row": 0, "max_column": 0, "has_header": False}

    # 提取所有行数据
    for row in worksheet.iter_rows(
        values_only=True, max_row=max_row, max_col=max_column
    ):
        # 过滤完全空的行
        if any(cell is not None and str(cell).strip() for cell in row):
            # 将单元格值转换为字符串，处理None值和空白
            row_content = []
            for cell in row:
                if cell is None:
                    row_content.append("")
                else:
                    row_content.append(str(cell).strip())
            rows.append("\t".join(row_content))

    # 检测是否有表头（第一行是否与其他行格式不同）
    has_header = _detect_header(worksheet) if rows else False

    return {
        "content": "\n".join(rows),
        "max_row": len(rows),
        "max_column": max_column,
        "has_header": has_header,
    }


class ExcelLoader(BaseLoader):
    def __init__(
        self,
        file_path: str | Path,
        mode: Literal["elements", "single"] = "single",
        **openpyxl_kwargs: Any,
    ):
        self.file_path = str(file_path)
        self.mode = mode
        self.openpyxl_kwargs = openpyxl_kwargs

        # 处理用户目录符号
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # 检查文件是否存在
        if not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file")

    def lazy_load(self) -> Iterator[Document]:
        elements = self._get_elements()

        if self.mode == "single":
            # 合并所有元素为单个文档
            combined_content = []
            combined_metadata = {
                "source": self.file_path,
                "mode": "single",
                "total_sheets": len(elements),
            }

            for element in elements:
                sheet_section = (
                    f"=== {element['metadata']['sheet_name']} ===\n{element['text']}"
                )
                combined_content.append(sheet_section)

            yield Document(
                page_content="\n\n".join(combined_content), metadata=combined_metadata
            )

        else:  # mode == "elements"
            # 每个工作表作为独立文档
            for element in elements:
                yield Document(
                    page_content=element["text"],
                    metadata={
                        "source": self.file_path,
                        "mode": "elements",
                        **element["metadata"],
                    },
                )

    def _get_elements(self) -> list[dict]:
        # 设置默认参数
        default_kwargs = {"read_only": True, "data_only": True}
        default_kwargs.update(self.openpyxl_kwargs)

        workbook = load_workbook(self.file_path, **default_kwargs)
        elements = []

        try:
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]

                # 提取工作表数据
                sheet_data = _extract_sheet_data(worksheet)

                if sheet_data["content"]:  # 只有非空工作表才添加
                    element = {
                        "type": "Table",
                        "text": sheet_data["content"],
                        "metadata": {
                            "sheet_name": sheet_name,
                            "filename": os.path.basename(self.file_path),
                            "max_row": sheet_data["max_row"],
                            "max_column": sheet_data["max_column"],
                            "has_header": sheet_data["has_header"],
                        },
                    }
                    elements.append(element)
        finally:
            workbook.close()

        return elements


__all__ = ["WordLoader", "ExcelLoader"]
